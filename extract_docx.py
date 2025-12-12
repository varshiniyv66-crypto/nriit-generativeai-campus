from docx import Document
import os

filesources_dir = "FILESOURCES"
docx_files = [
    "part-1 old.docx",
    "part-2 old.docx",
    "part-3 old.docx", 
    "part-4 old.docx",
    "part-5 old.docx"
]

for docx_file in docx_files:
    filepath = os.path.join(filesources_dir, docx_file)
    print(f"\n\n{'='*80}")
    print(f"FILE: {docx_file}")
    print('='*80)
    
    try:
        doc = Document(filepath)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                print(text)
        
        # Extract tables
        for table in doc.tables:
            print("\n--- TABLE ---")
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                if row_text.strip('| '):
                    print(row_text)
                    
    except Exception as e:
        print(f"Error reading {docx_file}: {e}")
